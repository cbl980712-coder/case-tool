# -*- coding: utf-8 -*-
"""
批量起诉状文书生成工具 v1
功能：解析名单 → 生成 Excel 总表 → 批量替换 Word 模板 → 生成异常报告
运行：python run_case_tool.py
"""

import os
import re
import sys
import copy
import pandas as pd
from pathlib import Path
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment
from docx import Document
from docx.oxml.ns import qn
import warnings
warnings.filterwarnings("ignore")

# ============================================================
# 路径配置
# ============================================================
BASE_DIR = Path(__file__).parent
INPUT_FILE = BASE_DIR / "input_list.txt"
TEMPLATE_FILE = BASE_DIR / "template.docx"
OUTPUT_EXCEL = BASE_DIR / "output" / "excel" / "case_list.xlsx"
OUTPUT_REPORT = BASE_DIR / "output" / "reports" / "error_report.xlsx"
OUTPUT_DOCS = BASE_DIR / "output" / "docs"

# 确保输出目录存在
for d in [OUTPUT_EXCEL.parent, OUTPUT_REPORT.parent, OUTPUT_DOCS]:
    d.mkdir(parents=True, exist_ok=True)

# ============================================================
# 第一版固定字段定义
# ============================================================
FIELDS = ["姓名", "证件号", "金额", "手机号", "地址", "备注"]
PLACEHOLDER_MAP = {
    "姓名":  "（姓名）",
    "证件号": "（证件号）",
    "金额":  "（金额）",
    "手机号": "（手机号）",
    "地址":  "（地址）",
    "备注":  "（备注）",
    "案号":  "（案号）",
}

# ============================================================
# 1. 解析名单文本
# ============================================================
def parse_input_list(filepath: Path) -> list[dict]:
    """读取 input_list.txt，返回人员信息列表。"""
    if not filepath.exists():
        print(f"[错误] 找不到输入文件：{filepath}")
        sys.exit(1)

    with open(filepath, "r", encoding="utf-8") as f:
        content = f.read()

    # 支持用空行分隔的块（允许块之间有任意空行）
    blocks = re.split(r"\n{2,}", content.strip())
    records = []
    for idx, block in enumerate(blocks, 1):
        block = block.strip()
        if not block:
            continue
        record = {"序号": idx, "案号": ""}
        for line in block.splitlines():
            line = line.strip()
            if "：" in line:
                key, _, val = line.partition("：")
                key = key.strip()
                val = val.strip()
                if key in FIELDS:
                    record[key] = val
        # 补全缺失字段
        for f in FIELDS:
            if f not in record:
                record[f] = ""
        records.append(record)

    print(f"[解析] 共读取到 {len(records)} 条记录")
    return records


# ============================================================
# 2. 数据校验，返回异常列表
# ============================================================
ID_PATTERN = re.compile(r"^\d{17}[\dXx]$")
PHONE_PATTERN = re.compile(r"^1[3-9]\d{9}$")
AMOUNT_PATTERN = re.compile(r"^\d+(\.\d{1,2})?$")

def validate_records(records: list[dict]) -> list[dict]:
    """对每条记录做格式校验，返回异常条目列表。"""
    errors = []
    seen_ids = {}

    for rec in records:
        row_errors = []
        seq = rec.get("序号", "?")
        name = rec.get("姓名", "").strip()
        id_no = rec.get("证件号", "").strip()
        amount = rec.get("金额", "").strip()
        phone = rec.get("手机号", "").strip()
        address = rec.get("地址", "").strip()

        if not name:
            row_errors.append("姓名为空")
        if not id_no:
            row_errors.append("证件号为空")
        elif not ID_PATTERN.match(id_no):
            row_errors.append("证件号格式异常（非18位）")
        else:
            if id_no in seen_ids:
                row_errors.append(f"证件号重复（与序号{seen_ids[id_no]}相同）")
            else:
                seen_ids[id_no] = seq

        if not amount:
            row_errors.append("金额为空")
        elif not AMOUNT_PATTERN.match(amount):
            row_errors.append("金额格式异常（非数字）")

        if not phone:
            row_errors.append("手机号为空")
        elif not PHONE_PATTERN.match(phone):
            row_errors.append("手机号格式异常（非11位大陆手机号）")

        if not address:
            row_errors.append("地址为空")

        if row_errors:
            errors.append({
                "序号": seq,
                "姓名": name,
                "证件号": id_no,
                "异常说明": "；".join(row_errors),
            })
            rec["状态"] = "异常：" + "；".join(row_errors)
        else:
            rec["状态"] = "正常"

    return errors


# ============================================================
# 3. 写入 Excel 总表
# ============================================================
def write_excel(records: list[dict], filepath: Path):
    """将解析结果写入 Excel 总表，保留案号列供后续补录。"""
    cols = ["序号", "姓名", "证件号", "金额", "手机号", "地址", "备注", "案号", "状态"]
    wb = Workbook()
    ws = wb.active
    ws.title = "案件总表"

    # 表头样式
    header_font = Font(bold=True, color="FFFFFF")
    header_fill = PatternFill("solid", fgColor="2F5496")
    for col_idx, col_name in enumerate(cols, 1):
        cell = ws.cell(row=1, column=col_idx, value=col_name)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = Alignment(horizontal="center")

    # 数据行
    for rec in records:
        row = [rec.get(c, "") for c in cols]
        ws.append(row)

    # 自动列宽（近似）
    for col in ws.columns:
        max_len = max(len(str(cell.value or "")) for cell in col)
        ws.column_dimensions[col[0].column_letter].width = min(max_len * 2.2, 60)

    wb.save(filepath)
    print(f"[Excel] 总表已保存：{filepath}")


# ============================================================
# 4. 写入异常报告
# ============================================================
def write_error_report(errors: list[dict], filepath: Path):
    """将校验异常写入独立的 Excel 报告。"""
    wb = Workbook()
    ws = wb.active
    ws.title = "异常报告"

    header_font = Font(bold=True, color="FFFFFF")
    header_fill = PatternFill("solid", fgColor="C00000")
    cols = ["序号", "姓名", "证件号", "异常说明"]
    for col_idx, col_name in enumerate(cols, 1):
        cell = ws.cell(row=1, column=col_idx, value=col_name)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = Alignment(horizontal="center")

    for err in errors:
        ws.append([err.get(c, "") for c in cols])

    for col in ws.columns:
        max_len = max(len(str(cell.value or "")) for cell in col)
        ws.column_dimensions[col[0].column_letter].width = min(max_len * 2.2, 80)

    wb.save(filepath)
    if errors:
        print(f"[报告] 发现 {len(errors)} 条异常，详见：{filepath}")
    else:
        print(f"[报告] 无异常，空报告已保存：{filepath}")


# ============================================================
# 5. Word 占位符替换工具函数
# ============================================================
def _replace_in_paragraph(para, replacements: dict):
    """
    在段落中替换占位符。
    由于 python-docx 可能将一个占位符拆分到多个 run 中，
    先合并段落全文做检测，再逐个 run 替换。
    """
    full_text = "".join(run.text for run in para.runs)
    for placeholder, value in replacements.items():
        if placeholder in full_text:
            full_text = full_text.replace(placeholder, value)
            # 清空所有 run，把全文写入第一个 run
            for i, run in enumerate(para.runs):
                run.text = full_text if i == 0 else ""
            break  # 单次替换后退出，re-enter 时 full_text 已更新
    # 多占位符时需要再次遍历，用递归方式简单处理
    full_text2 = "".join(run.text for run in para.runs)
    changed = False
    for placeholder, value in replacements.items():
        if placeholder in full_text2:
            full_text2 = full_text2.replace(placeholder, value)
            changed = True
    if changed:
        for i, run in enumerate(para.runs):
            run.text = full_text2 if i == 0 else ""


def _replace_in_table(table, replacements: dict):
    """遍历表格每个单元格每个段落进行替换。"""
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                _replace_in_paragraph(para, replacements)
            # 嵌套表格
            for nested_table in cell.tables:
                _replace_in_table(nested_table, replacements)


def replace_placeholders(doc: Document, replacements: dict) -> list[str]:
    """
    替换文档中所有段落和表格中的占位符。
    返回仍残留的占位符列表（用于警告）。
    """
    # 替换正文段落
    for para in doc.paragraphs:
        _replace_in_paragraph(para, replacements)

    # 替换表格
    for table in doc.tables:
        _replace_in_table(table, replacements)

    # 检测残留占位符
    remaining = []
    all_text = "\n".join(p.text for p in doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                all_text += "\n" + "\n".join(p.text for p in cell.paragraphs)
    # 匹配全角括号占位符
    found = re.findall(r"（[^）]{1,20}）", all_text)
    for f in found:
        if f not in remaining:
            remaining.append(f)
    return remaining


# ============================================================
# 6. 生成单个 Word 文书
# ============================================================
def generate_doc(rec: dict, template_path: Path, output_dir: Path) -> bool:
    """根据单条记录替换模板，生成文书。返回是否成功。"""
    name = rec.get("姓名", "未知")
    case_no = str(rec.get("案号", "")).strip()

    # 文件命名
    if case_no:
        filename = f"{name}_（{case_no}）_起诉状.docx"
    else:
        filename = f"{name}_起诉状.docx"

    # 过滤非法文件名字符
    filename = re.sub(r'[\\/:*?"<>|]', "_", filename)
    out_path = output_dir / filename

    # 构造替换字典
    replacements = {}
    for field, placeholder in PLACEHOLDER_MAP.items():
        replacements[placeholder] = str(rec.get(field, ""))

    try:
        doc = Document(str(template_path))
        remaining = replace_placeholders(doc, replacements)
        doc.save(str(out_path))

        if remaining:
            print(f"  [警告] {filename} 中仍残留占位符：{remaining}")
        return True
    except Exception as e:
        print(f"  [错误] 生成 {filename} 失败：{e}")
        return False


# ============================================================
# 7. 从 Excel 重新生成（补录案号后）
# ============================================================
def generate_from_excel(excel_path: Path, template_path: Path, output_dir: Path):
    """从已有的 Excel 总表读取数据并重新生成所有文书（支持补录案号）。"""
    if not excel_path.exists():
        print(f"[错误] 找不到 Excel 文件：{excel_path}")
        return

    df = pd.read_excel(str(excel_path), sheet_name="案件总表", dtype=str)
    df = df.fillna("")
    records = df.to_dict(orient="records")

    print(f"[模式] 从 Excel 重新生成文书，共 {len(records)} 条")
    success = 0
    for rec in records:
        # 跳过状态含"异常"的记录
        status = str(rec.get("状态", ""))
        if "异常" in status:
            print(f"  [跳过] {rec.get('姓名','')} — 有异常，不生成文书")
            continue
        if generate_doc(rec, template_path, output_dir):
            success += 1

    print(f"[完成] 从 Excel 重新生成：成功 {success}/{len(records)}")


# ============================================================
# 主流程
# ============================================================
def main():
    print("=" * 60)
    print("  批量起诉状文书生成工具 v1")
    print("=" * 60)

    # —— 模式判断 ——
    # 如果 Excel 总表已存在，询问用户是否从 Excel 重新生成
    if OUTPUT_EXCEL.exists():
        print(f"\n[检测] 发现已有 Excel 总表：{OUTPUT_EXCEL}")
        ans = input("是否从 Excel 重新生成文书（适用于已补录案号）？[y/N]: ").strip().lower()
        if ans == "y":
            if not TEMPLATE_FILE.exists():
                print(f"[错误] 找不到模板文件：{TEMPLATE_FILE}")
                sys.exit(1)
            generate_from_excel(OUTPUT_EXCEL, TEMPLATE_FILE, OUTPUT_DOCS)
            print("\n[结束] 重新生成完毕。")
            return

    # —— 正常流程：从名单文本开始 ——
    if not INPUT_FILE.exists():
        print(f"[错误] 找不到输入名单：{INPUT_FILE}")
        print("  请在脚本同级目录放置 input_list.txt")
        sys.exit(1)

    if not TEMPLATE_FILE.exists():
        print(f"[错误] 找不到模板文件：{TEMPLATE_FILE}")
        print("  请在脚本同级目录放置 template.docx")
        sys.exit(1)

    # Step 1: 解析
    records = parse_input_list(INPUT_FILE)

    # Step 2: 校验
    errors = validate_records(records)

    # Step 3: 写 Excel 总表
    write_excel(records, OUTPUT_EXCEL)

    # Step 4: 写异常报告
    write_error_report(errors, OUTPUT_REPORT)

    # Step 5: 批量生成文书（跳过有异常的记录）
    if not records:
        print("[提示] 没有可生成的记录，退出。")
        return

    print(f"\n[文书生成] 开始生成，共 {len(records)} 条……")
    success = 0
    skip = 0
    for rec in records:
        status = rec.get("状态", "")
        if "异常" in status:
            print(f"  [跳过] {rec.get('姓名','')} — 有异常：{status}")
            skip += 1
            continue
        if generate_doc(rec, TEMPLATE_FILE, OUTPUT_DOCS):
            success += 1
            print(f"  [生成] {rec.get('姓名','')} ✓")

    print(f"\n[汇总]")
    print(f"  总记录数  : {len(records)}")
    print(f"  生成成功  : {success}")
    print(f"  跳过（异常）: {skip}")
    print(f"  异常报告  : {OUTPUT_REPORT}")
    print(f"  Excel 总表 : {OUTPUT_EXCEL}")
    print(f"  文书输出目录: {OUTPUT_DOCS}")
    print("\n[提示] 如需补录案号，请在 Excel 总表的[案号]列填入内容后，")
    print("  重新运行本脚本，选择 [y] 从 Excel 重新生成。")
    print("\n[完成]")


if __name__ == "__main__":
    main()
