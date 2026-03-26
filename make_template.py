"""
生成 template.docx 示例模板（仅用于测试，实际由律师准备）
运行：python make_template.py
"""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# 标题
title = doc.add_heading("民事起诉状", level=1)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 正文
doc.add_paragraph("原告：（姓名），身份证号：（证件号），住址：（地址），联系电话：（手机号）。")
doc.add_paragraph("被告：某某培训机构（统一社会信用代码：XXXXXXXXXXXXXXXX）。")
doc.add_paragraph("案由：合同纠纷。")

# 诉讼请求标题
doc.add_heading("诉讼请求", level=2)
doc.add_paragraph("一、请求判令被告返还原告培训费（金额）元；")
doc.add_paragraph("二、请求判令被告承担本案全部诉讼费用。")

# 事实与理由
doc.add_heading("事实与理由", level=2)
doc.add_paragraph(
    "原告与被告于2024年签订培训服务合同，原告依约支付培训费（金额）元。"
    "被告在合同履行过程中无故停课，构成根本违约，导致原告遭受经济损失。"
    "原告多次协商无果，现依法提起诉讼，请求贵院支持原告诉讼请求。"
)
doc.add_paragraph("（备注）")

# 证据清单（表格示例）
doc.add_heading("证据清单", level=2)
table = doc.add_table(rows=1, cols=3)
table.style = "Table Grid"
hdr = table.rows[0].cells
hdr[0].text = "序号"
hdr[1].text = "证据名称"
hdr[2].text = "证明目的"
row = table.add_row().cells
row[0].text = "1"
row[1].text = "原告（姓名）身份证复印件"
row[2].text = "证明原告身份"
row2 = table.add_row().cells
row2[0].text = "2"
row2[1].text = "合同及转账凭证（金额元）"
row2[2].text = "证明合同关系及付款事实"

# 结尾
doc.add_paragraph("")
doc.add_paragraph("此致")
doc.add_paragraph("XX人民法院")
doc.add_paragraph("")
doc.add_paragraph("具状人：（姓名）")
doc.add_paragraph("日期：____年____月____日")

doc.save("template.docx")
print("template.docx 已生成")
