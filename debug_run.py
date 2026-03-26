# -*- coding: utf-8 -*-
"""
最小调试脚本 debug_run.py
不带GUI，直接测试核心生成链路
放在和 template.docx 同一目录运行：python debug_run.py
"""
import os, re, sys, traceback
from pathlib import Path

print("="*50)
print("核心链路调试脚本启动")
print("="*50)

# ── 步骤1：硬编码一条测试数据 ──
print("\n步骤1：准备测试数据")
record = {
    "姓名": "张三",
    "证件号": "350123199901011234",
    "金额": "12800",
    "手机号": "13800000000",
    "地址": "福建省福州市鼓楼区XX路XX号",
    "备注": "无",
    "案号": ""
}
print("测试数据：", record)

# ── 步骤2：找模板 ──
print("\n步骤2：查找模板文件")
tmpl = Path("template.docx")
if not tmpl.exists():
    print("❌ 找不到 template.docx，请把模板放在本脚本同目录")
    input("按回车退出...")
    sys.exit(1)
print("✅ 模板找到：", tmpl.resolve())

# ── 步骤3：打开模板 ──
print("\n步骤3：打开模板")
try:
    from docx import Document
    doc = Document(str(tmpl))
    print("✅ 模板打开成功")
except Exception as e:
    print("❌ 打开模板失败：", e)
    traceback.print_exc()
    input("按回车退出...")
    sys.exit(1)

# ── 步骤4：扫描占位符 ──
print("\n步骤4：扫描模板中的占位符")
all_text = "\n".join(p.text for p in doc.paragraphs)
for t in doc.tables:
    for row in t.rows:
        for cell in row.cells:
            all_text += "\n" + "\n".join(p.text for p in cell.paragraphs)
found = re.findall(r"（[^）]{1,20}）", all_text)
if found:
    print("✅ 检测到占位符：", list(dict.fromkeys(found)))
else:
    print("⚠️  未检测到任何占位符！请检查模板里用的是不是全角括号（姓名）")

# ── 步骤5：替换 ──
print("\n步骤5：替换占位符")
FIELDS = ["姓名","证件号","金额","手机号","地址","备注","案号"]
rmap = {f"（{f}）": str(record.get(f,"")) for f in FIELDS}
print("替换映射：", rmap)

def replace_para(para, rmap):
    full = "".join(r.text for r in para.runs)
    changed = False
    for ph, val in rmap.items():
        if ph in full:
            full = full.replace(ph, val)
            changed = True
    if changed:
        for i, run in enumerate(para.runs):
            run.text = full if i == 0 else ""

for para in doc.paragraphs:
    replace_para(para, rmap)
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                replace_para(para, rmap)
print("✅ 替换完成")

# ── 步骤6：保存 ──
print("\n步骤6：保存文档")
out = Path("调试输出_张三_起诉状.docx")
try:
    doc.save(str(out))
    print("✅ 保存成功：", out.resolve())
    print("\n请打开上面的文件检查内容是否正确")
except Exception as e:
    print("❌ 保存失败：", e)
    traceback.print_exc()

print("\n" + "="*50)
print("调试完成")
input("按回车退出...")
