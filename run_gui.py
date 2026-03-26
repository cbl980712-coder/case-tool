# -*- coding: utf-8 -*-
"""
批量文书生成工具 v1 - 桌面版
运行：python run_gui.py
依赖：pip install python-docx openpyxl pandas
"""

import os, re, sys, copy, threading, subprocess
import warnings
warnings.filterwarnings("ignore")

# ── 启动前检查依赖，缺少时弹窗提示而不是直接闪退 ──
def _check_deps():
    missing = []
    try:
        import tkinter
    except ImportError:
        # tkinter 缺失时只能打印，没有窗口
        print("错误：tkinter 未安装。Windows请重装Python并勾选tcl/tk，Mac请安装python-tk。")
        sys.exit(1)
    try:
        import pandas
    except ImportError:
        missing.append("pandas")
    try:
        import openpyxl
    except ImportError:
        missing.append("openpyxl")
    try:
        import docx
    except ImportError:
        missing.append("python-docx")
    if missing:
        import tkinter as _tk
        import tkinter.messagebox as _mb
        _r = _tk.Tk(); _r.withdraw()
        _mb.showerror(
            "缺少依赖",
            f"以下依赖未安装，请先在命令行运行：\n\n"
            f"pip install {' '.join(missing)}\n\n"
            f"安装完成后重新运行本程序。"
        )
        sys.exit(1)

_check_deps()

import tkinter as tk
from tkinter import ttk, filedialog, messagebox, scrolledtext
from pathlib import Path
import pandas as pd
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment
from docx import Document

# ─────────────────────────────────────────────
# 校验规则
# ─────────────────────────────────────────────
ID_PATTERN    = re.compile(r"^\d{17}[\dXx]$")
PHONE_PATTERN = re.compile(r"^1[3-9]\d{9}$")
AMOUNT_PATTERN= re.compile(r"^\d+(\.\d{1,2})?$")

PLACEHOLDER_PREFIX = "（"
PLACEHOLDER_SUFFIX = "）"

def make_placeholder(field):
    return PLACEHOLDER_PREFIX + field + PLACEHOLDER_SUFFIX

# ─────────────────────────────────────────────
# 解析名单文本
# ─────────────────────────────────────────────
def parse_text(text: str, fields: list) -> list:
    blocks = re.split(r"\n{2,}", text.strip())
    records = []
    for idx, block in enumerate(blocks, 1):
        block = block.strip()
        if not block:
            continue
        record = {"序号": idx, "案号": ""}
        for line in block.splitlines():
            line = line.strip()
            # 同时支持中文冒号"："和英文冒号":"
            sep = None
            if "：" in line:
                sep = "："
            elif ":" in line:
                sep = ":"
            if sep:
                key, _, val = line.partition(sep)
                key = key.strip(); val = val.strip()
                if key in fields:
                    record[key] = val
        for f in fields:
            if f not in record:
                record[f] = ""
        records.append(record)
    return records

# ─────────────────────────────────────────────
# 校验
# ─────────────────────────────────────────────
def validate_records(records, fields):
    errors = []
    seen_ids = {}
    required = {"姓名", "证件号", "金额", "手机号", "地址"}
    for rec in records:
        row_errors = []
        seq  = rec.get("序号", "?")
        name = rec.get("姓名", "").strip()
        id_no= rec.get("证件号","").strip()
        amt  = rec.get("金额","").strip()
        phone= rec.get("手机号","").strip()
        addr = rec.get("地址","").strip()

        if "姓名"  in fields and not name:  row_errors.append("姓名为空")
        if "证件号" in fields:
            if not id_no: row_errors.append("证件号为空")
            elif not ID_PATTERN.match(id_no): row_errors.append("证件号格式异常（非18位）")
            else:
                if id_no in seen_ids: row_errors.append(f"证件号重复（与序号{seen_ids[id_no]}相同）")
                else: seen_ids[id_no] = seq
        if "金额" in fields:
            if not amt: row_errors.append("金额为空")
            elif not AMOUNT_PATTERN.match(amt): row_errors.append("金额格式异常")
        if "手机号" in fields:
            if not phone: row_errors.append("手机号为空")
            elif not PHONE_PATTERN.match(phone): row_errors.append("手机号格式异常")
        if "地址" in fields and not addr: row_errors.append("地址为空")

        if row_errors:
            errors.append({"序号":seq,"姓名":name,"证件号":id_no,"异常说明":"；".join(row_errors)})
            rec["状态"] = "异常：" + "；".join(row_errors)
        else:
            rec["状态"] = "正常"
    return errors

# ─────────────────────────────────────────────
# 写 Excel
# ─────────────────────────────────────────────
def write_excel(records, fields, filepath):
    cols = ["序号"] + fields + ["案号","状态"]
    wb = Workbook(); ws = wb.active; ws.title="案件总表"
    hf = Font(bold=True, color="FFFFFF")
    hfill = PatternFill("solid", fgColor="2F5496")
    for ci, cn in enumerate(cols,1):
        c = ws.cell(row=1,column=ci,value=cn)
        c.font=hf; c.fill=hfill; c.alignment=Alignment(horizontal="center")
    for rec in records:
        ws.append([rec.get(c,"") for c in cols])
    for col in ws.columns:
        ml = max(len(str(cell.value or "")) for cell in col)
        ws.column_dimensions[col[0].column_letter].width = min(ml*2.2, 60)
    Path(filepath).parent.mkdir(parents=True, exist_ok=True)
    wb.save(str(filepath))

def write_error_report(errors, filepath):
    wb = Workbook(); ws = wb.active; ws.title="异常报告"
    hf = Font(bold=True,color="FFFFFF")
    hfill = PatternFill("solid",fgColor="C00000")
    cols=["序号","姓名","证件号","异常说明"]
    for ci,cn in enumerate(cols,1):
        c=ws.cell(row=1,column=ci,value=cn)
        c.font=hf; c.fill=hfill; c.alignment=Alignment(horizontal="center")
    for err in errors:
        ws.append([err.get(c,"") for c in cols])
    for col in ws.columns:
        ml=max(len(str(cell.value or "")) for cell in col)
        ws.column_dimensions[col[0].column_letter].width=min(ml*2.2,80)
    Path(filepath).parent.mkdir(parents=True, exist_ok=True)
    wb.save(str(filepath))

# ─────────────────────────────────────────────
# Word 替换核心
# ─────────────────────────────────────────────
def _replace_para(para, rmap):
    full = "".join(r.text for r in para.runs)
    changed = False
    for ph, val in rmap.items():
        if ph in full:
            full = full.replace(ph, val)
            changed = True
    if changed:
        for i, run in enumerate(para.runs):
            run.text = full if i == 0 else ""

def _replace_table(table, rmap):
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                _replace_para(para, rmap)
            for nt in cell.tables:
                _replace_table(nt, rmap)

def replace_placeholders(doc, rmap):
    for para in doc.paragraphs:
        _replace_para(para, rmap)
    for table in doc.tables:
        _replace_table(table, rmap)
    # 检测残留
    all_text = "\n".join(p.text for p in doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                all_text += "\n" + "\n".join(p.text for p in cell.paragraphs)
    return re.findall(r"（[^）]{1,20}）", all_text)

# ─────────────────────────────────────────────
# 格式转换：.doc/.wps/.odt → .docx (依赖 LibreOffice)
# ─────────────────────────────────────────────
def convert_to_docx(src_path: str, log_fn=None) -> str:
    """用 LibreOffice 把非 docx 格式转成 docx，返回转换后路径。"""
    src = Path(src_path)
    if src.suffix.lower() == ".docx":
        return src_path
    if src.suffix.lower() == ".txt":
        return src_path   # txt 单独处理

    # 检查 LibreOffice
    lo_cmds = ["libreoffice", "soffice",
               r"C:\Program Files\LibreOffice\program\soffice.exe",
               r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
               "/Applications/LibreOffice.app/Contents/MacOS/soffice"]
    lo_exe = None
    for cmd in lo_cmds:
        try:
            subprocess.run([cmd, "--version"], capture_output=True, timeout=5)
            lo_exe = cmd; break
        except Exception:
            continue

    if lo_exe is None:
        raise RuntimeError("未检测到 LibreOffice，无法转换 .doc/.wps 格式。\n"
                           "请安装 LibreOffice（免费）后重试，或直接使用 .docx 格式。")

    out_dir = src.parent
    result = subprocess.run(
        [lo_exe, "--headless", "--convert-to", "docx", "--outdir", str(out_dir), str(src)],
        capture_output=True, text=True, timeout=30
    )
    if result.returncode != 0:
        raise RuntimeError(f"LibreOffice 转换失败：{result.stderr}")

    converted = out_dir / (src.stem + ".docx")
    if not converted.exists():
        raise RuntimeError(f"转换后文件未找到：{converted}")
    if log_fn:
        log_fn(f"[转换] {src.name} → {converted.name}")
    return str(converted)

# ─────────────────────────────────────────────
# 生成单份文书（docx）
# ─────────────────────────────────────────────
def safe_filename(name):
    return re.sub(r'[\\/:*?"<>|]', "_", name)

def generate_one_doc(rec, fields, template_path, output_dir, log_fn=None):
    name    = rec.get("姓名","未知").strip() or "未知"
    case_no = str(rec.get("案号","")).strip()
    tmpl_p  = Path(template_path)
    suffix  = tmpl_p.suffix.lower()

    if case_no:
        base = f"{safe_filename(name)}_（{safe_filename(case_no)}）_起诉状"
    else:
        base = f"{safe_filename(name)}_起诉状"

    # 纯文本模板
    if suffix == ".txt":
        out_file = Path(output_dir) / (base + ".txt")
        txt = tmpl_p.read_text(encoding="utf-8")
        rmap = {make_placeholder(f): str(rec.get(f,"")) for f in fields}
        rmap[make_placeholder("案号")] = case_no
        for ph, val in rmap.items():
            txt = txt.replace(ph, val)
        out_file.parent.mkdir(parents=True, exist_ok=True)
        out_file.write_text(txt, encoding="utf-8")
        if log_fn: log_fn(f"  ✅ {out_file.name}")
        return True

    # docx / 其他（先转换）
    actual_tmpl = convert_to_docx(str(tmpl_p), log_fn)
    out_file = Path(output_dir) / (base + ".docx")
    out_file.parent.mkdir(parents=True, exist_ok=True)

    rmap = {make_placeholder(f): str(rec.get(f,"")) for f in fields}
    rmap[make_placeholder("案号")] = case_no

    doc = Document(actual_tmpl)
    remaining = replace_placeholders(doc, rmap)
    doc.save(str(out_file))
    if remaining and log_fn:
        log_fn(f"  ⚠️  {out_file.name} 残留占位符：{remaining}")
    elif log_fn:
        log_fn(f"  ✅ {out_file.name}")
    return True

# ─────────────────────────────────────────────
# GUI 主窗口
# ─────────────────────────────────────────────
class App(tk.Tk):
    def __init__(self):
        super().__init__()
        self.title("批量文书生成工具 v1")
        self.geometry("860x780")
        self.resizable(True, True)
        self.configure(bg="#F0F4F8")

        # 状态变量
        self.input_text_var   = tk.StringVar()
        self.template_path    = tk.StringVar()
        self.template_fmt     = tk.StringVar(value=".docx")
        self.excel_path       = tk.StringVar()
        self.output_dir       = tk.StringVar(value=str(Path.home() / "文书输出"))
        self.custom_fields    = []   # 用户自定义字段
        self._base_fields     = ["姓名","证件号","金额","手机号","地址","备注"]

        self._build_ui()

    # ── 构建 UI ──────────────────────────────
    def _build_ui(self):
        style = ttk.Style(self)
        style.theme_use("clam")
        style.configure("TLabelframe.Label", font=("Microsoft YaHei",10,"bold") if True else font, foreground="#2F5496")
        style.configure("Run.TButton", font=("Microsoft YaHei",11,"bold"),
                        background="#2F5496", foreground="white", padding=8)
        style.map("Run.TButton", background=[("active","#1a3a7a")])
        style.configure("Small.TButton", font=("Microsoft YaHei",9), padding=4)

        main = ttk.Frame(self, padding=12)
        main.pack(fill="both", expand=True)

        # ① 当事人信息
        f1 = ttk.LabelFrame(main, text="① 当事人信息输入", padding=8)
        f1.pack(fill="x", pady=(0,8))
        btn_row = ttk.Frame(f1)
        btn_row.pack(fill="x")
        ttk.Button(btn_row, text="📂 从文件导入 input_list.txt",
                   style="Small.TButton",
                   command=self._pick_input_file).pack(side="left", padx=(0,8))
        ttk.Label(btn_row, text="或直接在下方粘贴名单文本",
                  font=("Microsoft YaHei",9), foreground="#666").pack(side="left")
        ttk.Button(btn_row, text="清空",
                   style="Small.TButton",
                   command=lambda: self.input_box.delete("1.0","end")).pack(side="right")

        self.input_box = scrolledtext.ScrolledText(
            f1, height=10, font=("Consolas",10),
            wrap="word", relief="solid", bd=1)
        self.input_box.pack(fill="both", expand=True, pady=(6,0))

        # 自定义字段行
        cf_row = ttk.Frame(f1)
        cf_row.pack(fill="x", pady=(6,0))
        ttk.Label(cf_row, text="自定义字段：",
                  font=("Microsoft YaHei",9)).pack(side="left")
        self.custom_field_entry = ttk.Entry(cf_row, width=12)
        self.custom_field_entry.pack(side="left", padx=4)
        ttk.Button(cf_row, text="＋ 添加字段",
                   style="Small.TButton",
                   command=self._add_custom_field).pack(side="left")
        self.custom_field_label = ttk.Label(cf_row, text="",
                  font=("Microsoft YaHei",9), foreground="#2F5496")
        self.custom_field_label.pack(side="left", padx=8)

        # ② 模板文件
        f2 = ttk.LabelFrame(main, text="② 文档模板", padding=8)
        f2.pack(fill="x", pady=(0,8))
        tmpl_row = ttk.Frame(f2)
        tmpl_row.pack(fill="x")
        ttk.Button(tmpl_row, text="📂 选择模板文件",
                   style="Small.TButton",
                   command=self._pick_template).pack(side="left", padx=(0,8))
        ttk.Label(tmpl_row, textvariable=self.template_path,
                  font=("Microsoft YaHei",9), foreground="#333",
                  wraplength=560, anchor="w").pack(side="left", fill="x", expand=True)

        fmt_row = ttk.Frame(f2)
        fmt_row.pack(fill="x", pady=(6,0))
        ttk.Label(fmt_row, text="模板格式：",
                  font=("Microsoft YaHei",9)).pack(side="left")
        fmts = [
            (".docx", "Word (.docx)  推荐"),
            (".doc",  "旧版Word (.doc)  需LibreOffice"),
            (".wps",  "WPS文字 (.wps)  需LibreOffice"),
            (".odt",  "开放文档 (.odt)  需LibreOffice"),
            (".txt",  "纯文本 (.txt)"),
        ]
        for val, label in fmts:
            ttk.Radiobutton(fmt_row, text=label, variable=self.template_fmt,
                            value=val).pack(side="left", padx=6)

        # ③ 案号补录
        f3 = ttk.LabelFrame(main, text="③ 案号补录（可选，立案后使用）", padding=8)
        f3.pack(fill="x", pady=(0,8))
        cn_row = ttk.Frame(f3)
        cn_row.pack(fill="x")
        ttk.Button(cn_row, text="📂 导入已有 case_list.xlsx",
                   style="Small.TButton",
                   command=self._pick_excel).pack(side="left", padx=(0,8))
        ttk.Label(cn_row, textvariable=self.excel_path,
                  font=("Microsoft YaHei",9), foreground="#333",
                  wraplength=500, anchor="w").pack(side="left")
        ttk.Label(f3,
                  text="如未导入 Excel，则从名单文本生成（案号留空）。\n"
                       "补录案号后，导入 Excel 再次运行即可生成含案号版本。",
                  font=("Microsoft YaHei",9), foreground="#666").pack(anchor="w", pady=(4,0))

        # ④ 输出目录
        f4 = ttk.LabelFrame(main, text="④ 文书保存位置", padding=8)
        f4.pack(fill="x", pady=(0,8))
        od_row = ttk.Frame(f4)
        od_row.pack(fill="x")
        ttk.Button(od_row, text="📁 选择保存目录",
                   style="Small.TButton",
                   command=self._pick_output_dir).pack(side="left", padx=(0,8))
        ttk.Label(od_row, textvariable=self.output_dir,
                  font=("Microsoft YaHei",10,"bold") if True else font, foreground="#2F5496",
                  wraplength=560, anchor="w").pack(side="left", fill="x", expand=True)
        ttk.Label(f4,
                  text="生成的文书统一保存到此目录，文件名为对应当事人姓名（含案号时一并体现）。",
                  font=("Microsoft YaHei",9), foreground="#666").pack(anchor="w", pady=(4,0))

        # 操作按钮行
        btn_main = ttk.Frame(main)
        btn_main.pack(fill="x", pady=(4,8))
        self._run_btn = ttk.Button(btn_main, text="▶  开始生成文书",
                   style="Run.TButton",
                   command=self._run_threaded)
        self._run_btn.pack(side="left", padx=(0,12))
        ttk.Button(btn_main, text="📂 打开输出目录",
                   style="Small.TButton",
                   command=self._open_output_dir).pack(side="left", padx=(0,8))
        ttk.Button(btn_main, text="📊 打开异常报告",
                   style="Small.TButton",
                   command=self._open_report).pack(side="left")

        # 进度条
        self.progress = ttk.Progressbar(main, mode="determinate")
        self.progress.pack(fill="x", pady=(0,6))

        # 日志区
        log_frame = ttk.LabelFrame(main, text="运行日志", padding=6)
        log_frame.pack(fill="both", expand=True)
        self.log_box = scrolledtext.ScrolledText(
            log_frame, height=10, font=("Consolas",9),
            state="disabled", relief="solid", bd=1,
            bg="#1E1E1E", fg="#D4D4D4",
            insertbackground="white")
        self.log_box.pack(fill="both", expand=True)

        self._log("准备就绪。请填写信息后点击【开始生成文书】。")

    # ── 辅助交互 ─────────────────────────────
    def _add_custom_field(self):
        name = self.custom_field_entry.get().strip()
        if not name:
            return
        if name in self._base_fields or name in self.custom_fields:
            messagebox.showwarning("提示", f"字段「{name}」已存在")
            return
        self.custom_fields.append(name)
        self.custom_field_entry.delete(0, "end")
        self.custom_field_label.config(
            text="已添加：" + " | ".join(self.custom_fields))
        self._log(f"[字段] 已添加自定义字段：（{name}）")

    def _pick_input_file(self):
        p = filedialog.askopenfilename(
            title="选择名单文件", filetypes=[("文本文件","*.txt"),("所有文件","*.*")])
        if p:
            txt = Path(p).read_text(encoding="utf-8")
            self.input_box.delete("1.0","end")
            self.input_box.insert("1.0", txt)
            self._log(f"[导入] {p}")

    def _pick_template(self):
        p = filedialog.askopenfilename(
            title="选择模板文件",
            filetypes=[("文档文件","*.docx *.doc *.wps *.odt *.txt"),
                       ("Word文档","*.docx"),("旧版Word","*.doc"),
                       ("WPS文字","*.wps"),("文本文件","*.txt"),
                       ("所有文件","*.*")])
        if p:
            self.template_path.set(p)
            ext = Path(p).suffix.lower()
            self.template_fmt.set(ext if ext in [".docx",".doc",".wps",".odt",".txt"] else ".docx")
            self._log(f"[模板] {p}")
            self._detect_placeholders(p)

    def _detect_placeholders(self, path):
        """扫描模板占位符，给用户预览提示。"""
        try:
            ext = Path(path).suffix.lower()
            if ext == ".txt":
                txt = Path(path).read_text(encoding="utf-8")
                found = re.findall(r"（[^）]{1,20}）", txt)
            elif ext == ".docx":
                doc = Document(path)
                all_text = "\n".join(p.text for p in doc.paragraphs)
                for t in doc.tables:
                    for row in t.rows:
                        for cell in row.cells:
                            all_text += "\n" + "\n".join(p.text for p in cell.paragraphs)
                found = re.findall(r"（[^）]{1,20}）", all_text)
            else:
                return
            if found:
                unique = list(dict.fromkeys(found))
                self._log(f"[模板扫描] 检测到占位符：{' '.join(unique)}")
            else:
                self._log("[模板扫描] ⚠️  未检测到任何占位符，请确认模板格式正确")
        except Exception as e:
            self._log(f"[模板扫描] 跳过（{e}）")

    def _pick_excel(self):
        p = filedialog.askopenfilename(
            title="选择已有总表 Excel",
            filetypes=[("Excel文件","*.xlsx"),("所有文件","*.*")])
        if p:
            self.excel_path.set(p)
            self._log(f"[Excel] 已选择：{p}")

    def _pick_output_dir(self):
        d = filedialog.askdirectory(title="选择文书保存目录")
        if d:
            self.output_dir.set(d)
            self._log(f"[输出] 保存目录：{d}")

    def _open_output_dir(self):
        d = self.output_dir.get()
        if not d or not Path(d).exists():
            messagebox.showinfo("提示","输出目录不存在，请先生成文书。"); return
        self._open_path(d)

    def _open_report(self):
        rp = Path(self.output_dir.get()) / "_reports" / "error_report.xlsx"
        if not rp.exists():
            messagebox.showinfo("提示","异常报告尚未生成。"); return
        self._open_path(str(rp))

    def _open_path(self, path):
        import platform
        system = platform.system()
        if system == "Windows":
            os.startfile(path)
        elif system == "Darwin":
            subprocess.Popen(["open", path])
        else:
            subprocess.Popen(["xdg-open", path])

    def _log(self, msg):
        def _do():
            self.log_box.config(state="normal")
            self.log_box.insert("end", msg + "\n")
            self.log_box.see("end")
            self.log_box.config(state="disabled")
        self.after(0, _do)

    def _set_progress(self, val, max_val=100):
        self.after(0, lambda: self.progress.config(value=val, maximum=max_val))

    # ── 主流程（线程中运行）────────────────────
    def _run_threaded(self):
        # 禁用按钮，防止重复点击
        self._set_btn_state("running")
        t = threading.Thread(target=self._run_safe, daemon=True)
        t.start()

    def _run_safe(self):
        try:
            self._run()
        except Exception as e:
            import traceback
            self._log(f"[严重错误] {e}")
            self._log(traceback.format_exc())
            self._set_btn_state("normal")

    def _set_btn_state(self, state):
        def _do():
            if state == "running":
                self._run_btn.config(text="⏳  生成中，请稍候...", state="disabled")
            else:
                self._run_btn.config(text="▶  开始生成文书", state="normal")
        self.after(0, _do)

    def _run(self):
        self._log("\n" + "="*50)
        self._log("开始执行...")
        self._log(f"[调试] Python版本: {sys.version}")
        self._set_progress(0)

        out_dir    = Path(self.output_dir.get().strip())
        excel_dir  = out_dir / "_excel"
        report_dir = out_dir / "_reports"
        docs_dir   = out_dir

        out_dir.mkdir(parents=True, exist_ok=True)
        excel_dir.mkdir(parents=True, exist_ok=True)
        report_dir.mkdir(parents=True, exist_ok=True)

        fields = self._base_fields + self.custom_fields
        excel_out   = excel_dir  / "case_list.xlsx"
        report_out  = report_dir / "error_report.xlsx"

        # ── 决定数据来源 ──────────────────────
        excel_src = self.excel_path.get().strip()
        if excel_src and Path(excel_src).exists():
            # 从 Excel 重新生成
            self._log(f"[模式] 从 Excel 总表重新生成：{excel_src}")
            try:
                import pandas as pd
                df = pd.read_excel(excel_src, sheet_name="案件总表", dtype=str)
                df = df.fillna("")
                records = df.to_dict(orient="records")
                # 序号补全
                for i, r in enumerate(records, 1):
                    if "序号" not in r: r["序号"] = i
                self._log(f"[解析] 从 Excel 读取 {len(records)} 条记录")
            except Exception as e:
                self._log(f"[错误] 读取 Excel 失败：{e}")
                self._set_btn_state("normal")
                return
            errors = []
        else:
            # 从文本框解析
            raw = self.input_box.get("1.0","end").strip()
            if not raw:
                self._log("[错误] 名单为空，请输入或导入名单文本")
                self._set_btn_state("normal")
                return
            records = parse_text(raw, fields)
            if not records:
                self._log("[错误] 解析结果为空，请检查名单格式")
                self._set_btn_state("normal")
                return
            self._log(f"[解析] 共 {len(records)} 条记录")
            self._set_progress(10)

            # 校验
            errors = validate_records(records, fields)
            self._set_progress(20)

            # 写 Excel
            write_excel(records, fields, excel_out)
            self._log(f"[Excel] 总表已保存：{excel_out}")
            self._set_progress(30)

            # 写异常报告
            write_error_report(errors, report_out)
            if errors:
                self._log(f"[报告] {len(errors)} 条异常，详见：{report_out}")
                for e in errors:
                    self._log(f"  ⚠️  序号{e['序号']} {e['姓名']}：{e['异常说明']}")
            else:
                self._log("[报告] 无异常")
            self._set_progress(40)

        # ── 模板检查 ──────────────────────────
        tmpl = self.template_path.get().strip()
        if not tmpl or not Path(tmpl).exists():
            self._log("[错误] 未选择模板文件，请先选择 template.docx")
            self._set_btn_state("normal")
            return

        # ── 批量生成 ──────────────────────────
        total   = len(records)
        success = 0; skip = 0
        self._log(f"\n[文书生成] 开始生成 {total} 份……")

        for idx, rec in enumerate(records, 1):
            status = str(rec.get("状态",""))
            name   = rec.get("姓名","未知")
            if "异常" in status:
                self._log(f"  ⏭️  跳过：{name}（有异常）")
                skip += 1
            else:
                try:
                    generate_one_doc(rec, fields, tmpl, str(docs_dir), self._log)
                    success += 1
                except Exception as e:
                    self._log(f"  ❌ {name} 生成失败：{e}")
            prog = 40 + int(idx / total * 60)
            self._set_progress(prog, 100)

        # ── 汇总 ──────────────────────────────
        self._set_progress(100)
        self._log(f"\n{'='*50}")
        self._log(f"✅ 完成！")
        self._log(f"   总记录：{total} 条")
        self._log(f"   成功生成：{success} 份")
        self._log(f"   跳过（异常）：{skip} 份")
        self._log(f"   保存目录：{docs_dir}")
        if errors:
            self._log(f"   异常报告：{report_out}")
        self._set_btn_state("normal")


# ─────────────────────────────────────────────
# 入口
# ─────────────────────────────────────────────
if __name__ == "__main__":
    try:
        app = App()
        app.mainloop()
    except Exception as e:
        import traceback
        # 写错误日志到桌面，方便排查
        log_path = Path.home() / "Desktop" / "case_tool_error.log"
        try:
            log_path.write_text(traceback.format_exc(), encoding="utf-8")
        except Exception:
            log_path = Path.home() / "case_tool_error.log"
            log_path.write_text(traceback.format_exc(), encoding="utf-8")
        try:
            messagebox.showerror("启动失败",
                f"程序出现错误：\n{e}\n\n错误日志已保存到：\n{log_path}")
        except Exception:
            print(traceback.format_exc())
